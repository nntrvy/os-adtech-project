#!/usr/bin/env python3
"""
Vietnam DOOH Advertising Analysis - Word Document Generator
Creates professionally formatted .docx report with embedded visualizations
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# File paths
OUTPUT_DIR = '/Users/nntrvy/os-adtech-project/'
OUTPUT_FILE = f'{OUTPUT_DIR}vietnam-dooh-advertiser-spending-analysis-visualized.docx'

def set_document_formatting(doc):
    """Configure document-wide formatting standards"""

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Set paragraph formatting
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15

    # Configure heading styles
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Arial'
    heading1.font.size = Pt(14)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)

    heading2 = doc.styles['Heading 2']
    heading2.font.name = 'Arial'
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)

    heading3 = doc.styles['Heading 3']
    heading3.font.name = 'Arial'
    heading3.font.size = Pt(11)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(0, 0, 0)

def add_title_page(doc):
    """Create title page"""

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Vietnam DOOH Advertiser Spending Analysis')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Who\'s Paying the Most Money into Digital Out-of-Home Advertising')
    run.font.size = Pt(14)
    run.font.italic = True

    doc.add_paragraph()  # Spacing

    # Report details
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run('Report Date: December 2, 2025\n').bold = True
    details.add_run('Market Focus: Vietnam DOOH Advertising Buyers & Spending Patterns\n')
    details.add_run('Time Period: 2024-2025 Market Intelligence')

    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing

    # Key metrics box
    metrics = doc.add_paragraph()
    metrics.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = metrics.add_run('KEY MARKET METRICS')
    run.font.size = Pt(12)
    run.font.bold = True

    metrics_text = doc.add_paragraph()
    metrics_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
    metrics_text.add_run('Total Market Size (2025): ').bold = True
    metrics_text.add_run('USD $129.58 Million\n')
    metrics_text.add_run('Market Growth Rate: ').bold = True
    metrics_text.add_run('5.74% - 12.54% CAGR\n')
    metrics_text.add_run('Top Spending Category: ').bold = True
    metrics_text.add_run('Retail & Consumer Goods (44.49%)\n')
    metrics_text.add_run('Fastest Growing Sector: ').bold = True
    metrics_text.add_run('Telecom & Technology (7.04% CAGR)\n')
    metrics_text.add_run('Geographic Concentration: ').bold = True
    metrics_text.add_run('HCMC + Hanoi (60-70% of spend)')

    doc.add_page_break()

def add_executive_summary(doc):
    """Add executive summary section"""

    doc.add_heading('Executive Summary', 1)

    intro = doc.add_paragraph(
        'This report provides detailed intelligence on advertiser spending in Vietnam\'s Digital Out-of-Home (DOOH) '
        'advertising market, identifying top spenders, budget allocations, and buyer behavior patterns.'
    )

    doc.add_heading('Key Findings', 2)

    findings = [
        'Retail & Consumer Goods dominate with 44.49% market share',
        'FMCG commands 33.71% of digital advertising spend',
        '91 of Vietnam\'s top 100 companies actively invest in DOOH',
        'Total market size: USD $129.58M (2025) growing to $171.29M (2030)',
        'Geographic concentration: HCMC and Hanoi capture majority of spending',
        'Digital DOOH growing at 12.54% CAGR, significantly outpacing static formats',
        'Programmatic adoption driving 61.8% revenue uplift for major operators'
    ]

    for finding in findings:
        p = doc.add_paragraph(finding, style='List Bullet')

def add_visualization_section_1(doc):
    """Industry spending and market share analysis"""

    doc.add_heading('1. Industry Spending Analysis', 1)

    doc.add_paragraph(
        'The Vietnam DOOH market shows clear industry concentration, with Retail and FMCG sectors '
        'accounting for nearly 78% of total advertising spending. This dominance reflects the strategic '
        'importance of physical presence and point-of-sale marketing in Vietnam\'s evolving retail landscape.'
    )

    # Chart 1: Market Share
    doc.add_heading('Figure 1: Industry Market Share Distribution', 2)

    if os.path.exists(f'{OUTPUT_DIR}chart1_market_share.png'):
        doc.add_picture(f'{OUTPUT_DIR}chart1_market_share.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph('Figure 1: Industry spending distribution showing Retail & Consumer Goods dominance at 44.49%, '
                                'followed by FMCG at 33.71%. These two sectors represent the majority of DOOH investment.')
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True

    # Key insights
    doc.add_heading('Key Insights: Industry Concentration', 3)

    insights = [
        'Retail & Consumer Goods (44.49%): Largest spender, driven by e-commerce competition and need for physical store traffic',
        'FMCG (33.71%): Heavy programmatic adoption enabling same-day creative optimization and audience retargeting',
        'Automotive (12%): Recording fastest budget increases, fueled by VinFast\'s domestic EV surge',
        'Combined Retail + FMCG: Control 78.2% of market, creating significant volume buying power'
    ]

    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')

    # Chart 7: Budget Allocation
    doc.add_heading('Figure 2: Annual Budget Allocation by Industry', 2)

    if os.path.exists(f'{OUTPUT_DIR}chart7_budget_allocation.png'):
        doc.add_picture(f'{OUTPUT_DIR}chart7_budget_allocation.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph('Figure 2: Absolute spending by industry in USD millions, illustrating the significant '
                                'investment gap between top-tier and emerging sectors.')
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True

def add_visualization_section_2(doc):
    """Growth rates and projections"""

    doc.add_heading('2. Growth Dynamics and Market Projections', 1)

    doc.add_paragraph(
        'Vietnam\'s DOOH market demonstrates strong growth momentum across all segments, with digital formats '
        'significantly outpacing traditional static billboards. The market is experiencing a fundamental shift '
        'toward programmatic buying and data-driven optimization.'
    )

    # Chart 2: Growth Rates
    doc.add_heading('Figure 3: Sector Growth Rates Comparison', 2)

    if os.path.exists(f'{OUTPUT_DIR}chart2_growth_rates.png'):
        doc.add_picture(f'{OUTPUT_DIR}chart2_growth_rates.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph('Figure 3: CAGR comparison showing Digital DOOH leading at 12.54%, followed by Mall/Retail '
                                'Interior and Telecom sectors. All segments outperform the total market average.')
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True

    # Key insights
    doc.add_heading('Growth Rate Analysis', 3)

    insights = [
        'Digital DOOH: 12.54% CAGR represents fastest growth trajectory, driven by programmatic capabilities',
        'Telecom & Technology: 7.04% CAGR fueled by 5G rollout and competitive wars among major operators',
        'Mall/Retail Interior: 7.70% CAGR reflects growing importance of point-of-sale proximity marketing',
        'Growth Differential: Digital formats growing 2.5x faster than static, signaling market transformation'
    ]

    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')

    # Chart 3: Market Projections
    doc.add_heading('Figure 4: Market Size Projections (2024-2030)', 2)

    if os.path.exists(f'{OUTPUT_DIR}chart3_market_projections.png'):
        doc.add_picture(f'{OUTPUT_DIR}chart3_market_projections.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph('Figure 4: Three projection scenarios showing market evolution. Digital DOOH (12.54% CAGR) '
                                'demonstrates steepest growth curve, while total market includes both static and digital segments.')
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True

    # Projections insight
    doc.add_heading('Market Evolution 2024-2030', 3)

    doc.add_paragraph(
        'Conservative projections suggest the market will reach $171M by 2030, while more aggressive '
        'digital-focused projections indicate potential for $282M if digital adoption accelerates. '
        'The most likely scenario places 2030 market size between $200-220M, representing a doubling '
        'of current market value.'
    )

def add_visualization_section_3(doc):
    """Geographic distribution"""

    doc.add_heading('3. Geographic Spending Patterns', 1)

    doc.add_paragraph(
        'DOOH spending in Vietnam shows pronounced geographic concentration, with Ho Chi Minh City and Hanoi '
        'capturing the majority of investment. This concentration reflects infrastructure advantages, '
        'population density, and advertiser targeting strategies focused on urban affluent consumers.'
    )

    # Chart 4: Geographic Distribution
    doc.add_heading('Figure 5: Geographic Spending Distribution', 2)

    if os.path.exists(f'{OUTPUT_DIR}chart4_geographic_distribution.png'):
        doc.add_picture(f'{OUTPUT_DIR}chart4_geographic_distribution.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph('Figure 5: Left panel shows spending concentration in Tier 1 cities (HCMC + Hanoi) at 65%. '
                                'Right panel compares key metrics between the two primary markets.')
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True

    # Geographic insights
    doc.add_heading('HCMC vs. Hanoi: Comparative Analysis', 3)

    comparison_text = [
        'Spending Volume: HCMC generates higher absolute spending due to larger population (9.4M vs 8.0M) and more developed infrastructure',
        'Digital Consumption: Hanoi residents spend MORE time online (3.1 hrs/day vs 2.6 hrs/day), showing stronger digital format preference',
        'Infrastructure: HCMC metro expansion creating hundreds of interior screens; Hanoi\'s smart city mandates driving municipal building displays',
        'Advertiser Focus: HCMC favors retail/FMCG campaigns; Hanoi over-indexes on tech/fintech programmatic buying'
    ]

    for text in comparison_text:
        doc.add_paragraph(text, style='List Bullet')

    doc.add_heading('Secondary Cities: Emerging Opportunities', 3)

    doc.add_paragraph(
        'Da Nang and other Tier 2 cities represent 30-40% of total spending, with lower CPMs enabling '
        'volume-based strategies. Static formats retain stronger presence in these markets, though '
        'urbanization (6.29% CAGR) is driving gradual digital adoption.'
    )

def add_visualization_section_4(doc):
    """Format breakdown and evolution"""

    doc.add_heading('4. Format Preferences: Static vs. Digital', 1)

    doc.add_paragraph(
        'Vietnam\'s DOOH market remains predominantly static (75.24% share) but is experiencing rapid '
        'digital transformation. Digital formats offer programmatic capabilities, real-time optimization, '
        'and superior attribution, driving accelerated adoption among sophisticated advertisers.'
    )

    # Chart 5: Format Breakdown
    doc.add_heading('Figure 6: Static vs. Digital Format Analysis', 2)

    if os.path.exists(f'{OUTPUT_DIR}chart5_format_breakdown.png'):
        doc.add_picture(f'{OUTPUT_DIR}chart5_format_breakdown.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph('Figure 6: Left panel shows current market share (2024) with static formats dominant. '
                                'Right panel projects spending evolution, highlighting digital\'s 2.5x faster growth rate.')
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True

    # Format insights
    doc.add_heading('Format Transition Dynamics', 3)

    insights = [
        'Current State: Static OOH maintains 75.24% share, concentrated in highway corridors and suburban locations',
        'Digital Growth: 12.54% CAGR represents 2.5x faster growth than static, driven by urban concentration and programmatic adoption',
        'Programmatic Impact: JCDecaux reported 61.8% platform revenue uplift via VIOOH SSP, demonstrating ROI of digital investment',
        'Investment Shift: Advertisers increasingly gravitating toward "data-rich digital screens" in HCMC and Hanoi'
    ]

    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')

    doc.add_heading('Format Preferences by Advertiser Type', 3)

    format_prefs = [
        'FMCG Brands: Programmatic-enabled screens for day-of creative optimization and audience retargeting',
        'Automotive: Highway megaboards (static and digital) for large-format brand building and test-drive promotions',
        'Retail: Mall interior screens near point-of-sale; complemented by mass-reach roadside billboards',
        'Telecom/Tech: Interactive metro screens driving app downloads and mobile actions',
        'Luxury Brands: Premium inventory with sophisticated 3D/CGI capabilities in affluent areas'
    ]

    for pref in format_prefs:
        doc.add_paragraph(pref, style='List Bullet')

def add_visualization_section_5(doc):
    """Seasonal spending patterns"""

    doc.add_heading('5. Seasonal Spending Patterns and Evolution', 1)

    doc.add_paragraph(
        'Vietnam\'s DOOH spending historically concentrated around Tet (Lunar New Year), with spirits brands '
        'generating 60% of yearly sales in the 6-week period. However, 2024 marked a significant shift, '
        'with traditional Tet campaigns showing 47% year-over-year decline in consumer engagement.'
    )

    # Chart 6: Seasonal Patterns
    doc.add_heading('Figure 7: Seasonal Spending Pattern Evolution', 2)

    if os.path.exists(f'{OUTPUT_DIR}chart6_seasonal_spending.png'):
        doc.add_picture(f'{OUTPUT_DIR}chart6_seasonal_spending.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph('Figure 7: Comparison of traditional Tet-concentrated spending vs. modern year-round optimization. '
                                'The 2024 shift reflects 47% engagement decline in traditional Tet campaigns.')
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True

    # Seasonal insights
    doc.add_heading('Strategic Shift: From Seasonal to Continuous', 3)

    insights = [
        'Traditional Peak: Tet period historically drove 200+ relative spending index, with spirits/beer dominating',
        '2024 Disruption: 47% YoY decline in consumer engagement; family reunion discussions down 77.5% on social media',
        'Modern Approach: Shift to year-round programmatic optimization with real-time creative adaptation',
        'New Drivers: Weather patterns, local events, traffic patterns, competitive activity, and sales data'
    ]

    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')

    doc.add_heading('Category-Specific Seasonal Patterns', 3)

    seasonal = [
        'FMCG: Weather-driven peaks (beverages in hot season), holiday promotional calendars, product launch windows',
        'Automotive: Model year launches (Q4-Q1), test-drive season coordination, auto show tie-ins',
        'Retail: Back-to-school, Black Friday (growing adoption), year-end shopping, e-commerce coordination',
        'Entertainment: Film release calendar synchronization, event-based shorter cycles, tactical high-impact bursts'
    ]

    for item in seasonal:
        doc.add_paragraph(item, style='List Bullet')

def add_key_advertisers_section(doc):
    """Major advertisers and campaign strategies"""

    doc.add_heading('6. Major Advertisers and Campaign Strategies', 1)

    doc.add_paragraph(
        '91 of Vietnam\'s top 100 companies by revenue actively invest in DOOH, with over 700 major domestic '
        'and foreign brands using the channel. This elite adoption demonstrates DOOH\'s strategic importance '
        'for reaching Vietnam\'s urban affluent consumers.'
    )

    doc.add_heading('Confirmed Heavy DOOH Spenders', 2)

    # Consumer Electronics
    doc.add_heading('Consumer Electronics', 3)
    doc.add_paragraph(
        'OPPO: Ran large-scale elevator advertising campaigns achieving breakthrough success surpassing Samsung '
        'in market share. Strategic DOOH deployment credited for market leadership shift.'
    )

    # FMCG
    doc.add_heading('FMCG & Beverage', 3)
    advertisers = [
        'Suntory PepsiCo Vietnam: Active Chicilon Media client with creative elevator media campaigns and programmatic adoption',
        'Nestle: Major participant in Vietnam DOOH ecosystem, featured in MMA Global forums, multi-format presence',
        'Heineken: Integrated billboard advertising with Tet seasonal campaigns, +30% beer sales driving seasonal investment'
    ]
    for adv in advertisers:
        doc.add_paragraph(adv, style='List Bullet')

    # Automotive
    doc.add_heading('Automotive', 3)
    doc.add_paragraph(
        'VinFast: Nationwide test-drive promotions via highway megaboards, historic Times Square presence '
        '(1,000 m² LED display), domestic EV surge fueling massive DOOH investment.'
    )

    # Luxury
    doc.add_heading('Luxury & Fashion', 3)
    doc.add_paragraph(
        'Calvin Klein: 27-day programmatic DOOH campaign with Moving Walls across 230 screens at 39 sites, '
        'generating 1.56 million billable impressions and 713K ad plays. Demonstrates luxury brand adoption '
        'of programmatic DOOH capabilities.'
    )

    # Campaign strategies by type
    doc.add_heading('Campaign Strategies by Advertiser Type', 2)

    doc.add_heading('FMCG Brands', 3)
    fmcg_strat = [
        'Tactical Approach: Programmatic buying for agility and day-of creative optimization',
        'Budget Allocation: Year-round presence with seasonal peaks, programmatic efficiency driving ROI',
        'Format Focus: Point-of-sale proximity targeting in malls, brand-safe environments for launches'
    ]
    for strat in fmcg_strat:
        doc.add_paragraph(strat, style='List Bullet')

    doc.add_heading('Luxury Brands', 3)
    luxury_strat = [
        'Tactical Approach: Data-driven location targeting, premium inventory in affluent areas, sophisticated 3D/CGI creative',
        'Budget Characteristics: Willing to pay premium CPMs, longer campaign durations, international coordination'
    ]
    for strat in luxury_strat:
        doc.add_paragraph(strat, style='List Bullet')

    doc.add_heading('Automotive', 3)
    auto_strat = [
        'Tactical Approach: Nationwide highway megaboard takeovers, test-drive event coordination, geographic conquest',
        'Budget Characteristics: Large format investments, model launch spikes, sustained brand presence'
    ]
    for strat in auto_strat:
        doc.add_paragraph(strat, style='List Bullet')

def add_market_drivers_section(doc):
    """Key growth drivers and opportunities"""

    doc.add_heading('7. Key Growth Drivers and Market Dynamics', 1)

    doc.add_heading('Primary Growth Catalysts', 2)

    drivers = [
        'Urbanization: Rapid metro expansion in HCMC and Hanoi creating new premium inventory in transit environments',
        'Smart City Infrastructure: Municipal buildings mandating digital signage, government-driven digital transformation',
        'Programmatic Adoption: Platforms (ANTS, VIOOH/JCDecaux, Moving Walls) enabling efficiency and real-time buying',
        '5G Rollout: Telecom competition fueling ad spend, enhanced interactive capabilities, mobile-to-DOOH integration',
        'E-commerce Pressure: Retailers fighting back with physical store traffic drivers, omnichannel strategies',
        'Measurement Improvement: Data-rich screens enabling attribution, programmatic analytics, improved ROI demonstration'
    ]

    for driver in drivers:
        doc.add_paragraph(driver, style='List Bullet')

    doc.add_heading('DOOH Unique Advantages', 2)

    advantages = [
        'Cannot be blocked (vs. digital ad blockers)',
        'Contextual relevance during commute/shopping moments',
        'Premium attention in urban environments',
        'Programmatic capabilities bridging digital and physical',
        'Brand safety in controlled environments',
        'Large format impact impossible in other channels'
    ]

    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')

    doc.add_heading('Market Challenges', 2)

    challenges = [
        'Tet Effectiveness Declining: 47% engagement drop requiring creative innovation beyond traditional messaging',
        'Measurement Standards Gap: Limited industry-standard metrics, attribution capabilities vary by vendor',
        'Programmatic Adoption Pace: Ongoing transition from traditional buying, learning curve for stakeholders',
        'Static to Digital Transition: 75% still static creates inventory risk, capital intensive to convert'
    ]

    for challenge in challenges:
        doc.add_paragraph(challenge, style='List Bullet')

def add_actionable_intelligence_section(doc):
    """Actionable insights for stakeholders"""

    doc.add_heading('8. Actionable Intelligence for Stakeholders', 1)

    doc.add_heading('For Media Buyers', 2)

    buyer_insights = [
        'Category Leverage: Retail/FMCG representing 78% of market creates volume discount opportunities',
        'Programmatic Access Required: Demand ANTS, VIOOH, or Moving Walls connectivity for optimization',
        'Seasonal Strategy Shift: Don\'t over-concentrate on Tet; year-round optimization preferred',
        'Geographic Arbitrage: Balance secondary city reach (lower CPMs) vs. Tier 1 precision (attribution data)',
        'Measurement Requirements: Demand attribution data, QR integration, store visit tracking, mobile retargeting'
    ]

    for insight in buyer_insights:
        doc.add_paragraph(insight, style='List Bullet')

    doc.add_heading('For Media Sellers', 2)

    seller_insights = [
        'Retail Sales Strategy: 44.49% share demands dedicated retail teams with POS attribution capabilities',
        'Programmatic Imperative: JCDecaux\'s 61.8% revenue uplift proves ROI; SSP connectivity critical',
        'Geographic Focus: HCMC + Hanoi = majority revenue; premium pricing justified in Tier 1 cities',
        'Format Transition: Digital growing 2.5x faster but don\'t abandon static too quickly; gradual conversion strategy'
    ]

    for insight in seller_insights:
        doc.add_paragraph(insight, style='List Bullet')

    doc.add_heading('For Investors/Strategists', 2)

    investor_insights = [
        'Growth Vector Priority: Digital DOOH at 12.54% CAGR; technology/platform plays preferred over traditional inventory',
        'Consolidation Opportunity: Fragmented market creates M&A potential for scale building',
        'Telecom Surge Preparation: 7.04% CAGR highest among sectors; 5G wars = ad spending wars',
        'Measurement Gap = Opportunity: Limited attribution creates opportunity for measurement tech platforms',
        'Retail Partnership Strategy: 44.49% market share makes retail media network co-development priority'
    ]

    for insight in investor_insights:
        doc.add_paragraph(insight, style='List Bullet')

def add_methodology_section(doc):
    """Data sources and limitations"""

    doc.add_heading('9. Methodology and Data Reliability', 1)

    doc.add_heading('Data Sources', 2)

    doc.add_paragraph(
        'This analysis synthesizes data from multiple market research firms (Mordor Intelligence, Statista, '
        'IMARC Group, Ken Research), industry case studies, company reports, and campaign data. Research '
        'reflects 2024 actual data and 2025 projections based on market intelligence through December 2024.'
    )

    doc.add_heading('Source Discrepancies and Limitations', 2)

    limitations = [
        'Market Size Range: Total Vietnam DOOH estimates vary from $32.6M to $500M due to different methodologies (digital-only vs. total OOH, geographic scope, format classification)',
        'Growth Projections: Conservative projections show 5.74% CAGR; aggressive projections show 10.14% CAGR depending on market definition',
        'Proprietary Data Gaps: Specific brand-level spending, detailed CPM/rate cards, programmatic platform market shares, and campaign-level ROI metrics not publicly available'
    ]

    for limitation in limitations:
        doc.add_paragraph(limitation, style='List Bullet')

    doc.add_heading('Recommended Further Research', 2)

    doc.add_paragraph(
        'For deeper insights, direct engagement recommended with: JCDecaux Vietnam, Chicilon Media, VCCorp, '
        'GroupM Vietnam, Dentsu Vietnam, Nielsen Vietnam ad spend tracking, and Vietnam Advertising Association reports.'
    )

def add_conclusion_section(doc):
    """Final conclusions and outlook"""

    doc.add_heading('10. Conclusions and Market Outlook', 1)

    doc.add_paragraph(
        'Vietnam\'s DOOH market stands at a transformative inflection point. With 91 of the country\'s top 100 '
        'companies actively investing, elite adoption is established. The market is experiencing three simultaneous '
        'shifts that will define its trajectory through 2030:'
    )

    conclusions = [
        'Digital Acceleration: 12.54% CAGR for digital formats vs. slower static growth signals inevitable format transition. Programmatic capabilities (61.8% revenue uplift for JCDecaux) demonstrate clear ROI advantage.',

        'Geographic Concentration with Secondary Growth: HCMC and Hanoi capture 60-70% of spending but secondary cities growing at 6.29% CAGR create geographic diversification opportunities.',

        'Seasonal to Continuous: The 47% Tet engagement decline marks the end of seasonal "spray and pray" strategies. Year-round programmatic optimization with real-time adaptation is the new standard.',

        'Category Evolution: While Retail (44.49%) and FMCG (33.71%) dominate today, Telecom\'s 7.04% CAGR signals emerging category shifts driven by 5G competition.',

        'Infrastructure Expansion: Metro development, smart city mandates, and airport LED installations are creating premium inventory that commands higher CPMs and attracts sophisticated programmatic buyers.'
    ]

    for conclusion in conclusions:
        p = doc.add_paragraph()
        p.add_run(conclusion.split(':')[0] + ':').bold = True
        p.add_run(' ' + ':'.join(conclusion.split(':')[1:]))

    doc.add_paragraph()
    doc.add_paragraph(
        'The market will likely reach $200-220M by 2030, representing a near-doubling from current levels. '
        'Success will favor players who embrace programmatic technology, develop measurement capabilities, '
        'cultivate retail partnerships, and maintain presence in both Tier 1 premium locations and emerging '
        'secondary cities.'
    )

def add_appendix(doc):
    """Appendix with data tables"""

    doc.add_page_break()
    doc.add_heading('Appendix: Key Data Tables', 1)

    doc.add_heading('Market Size Data', 2)

    # Create market size table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = '2024'
    header_cells[2].text = '2025'

    # Data rows
    data = [
        ['Digital OOH', '$107.50M', '$129.58M'],
        ['Static OOH', '$376M (75.24%)', '$395M (est.)'],
        ['Total Market', '$500M', '$530M (est.)']
    ]

    for i, row_data in enumerate(data):
        cells = table.rows[i+1].cells
        for j, value in enumerate(row_data):
            cells[j].text = value

    doc.add_paragraph()

    doc.add_heading('Industry Spending Breakdown (2025)', 2)

    # Create industry table
    table2 = doc.add_table(rows=7, cols=3)
    table2.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table2.rows[0].cells
    header_cells[0].text = 'Industry'
    header_cells[1].text = 'Market Share'
    header_cells[2].text = 'Annual Spending (USD)'

    # Data rows
    industry_data = [
        ['Retail & Consumer Goods', '44.49%', '$57.6M'],
        ['FMCG', '33.71%', '$43.7M'],
        ['Automotive', '~12%', '$15.5M'],
        ['BFSI', '~4.5%', '$5.8M'],
        ['Telecom & Technology', '1.93%', '$2.5M'],
        ['Others', '~3.37%', '$4.4M']
    ]

    for i, row_data in enumerate(industry_data):
        cells = table2.rows[i+1].cells
        for j, value in enumerate(row_data):
            cells[j].text = value

def create_report():
    """Main function to create the complete report"""

    print("\n" + "="*60)
    print("Vietnam DOOH Analysis - Word Document Generator")
    print("="*60 + "\n")

    # Create document
    doc = Document()

    # Set up formatting
    print("Setting document formatting...")
    set_document_formatting(doc)

    # Add sections
    print("Creating title page...")
    add_title_page(doc)

    print("Adding executive summary...")
    add_executive_summary(doc)

    print("Adding visualization section 1: Industry spending...")
    add_visualization_section_1(doc)

    print("Adding visualization section 2: Growth rates...")
    add_visualization_section_2(doc)

    print("Adding visualization section 3: Geographic distribution...")
    add_visualization_section_3(doc)

    print("Adding visualization section 4: Format breakdown...")
    add_visualization_section_4(doc)

    print("Adding visualization section 5: Seasonal patterns...")
    add_visualization_section_5(doc)

    print("Adding key advertisers section...")
    add_key_advertisers_section(doc)

    print("Adding market drivers section...")
    add_market_drivers_section(doc)

    print("Adding actionable intelligence...")
    add_actionable_intelligence_section(doc)

    print("Adding methodology section...")
    add_methodology_section(doc)

    print("Adding conclusions...")
    add_conclusion_section(doc)

    print("Adding appendix...")
    add_appendix(doc)

    # Save document
    print(f"\nSaving document to: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)

    print("\n" + "="*60)
    print("✓ Document created successfully!")
    print(f"Output file: {OUTPUT_FILE}")
    print("="*60 + "\n")

    return OUTPUT_FILE

if __name__ == "__main__":
    output_path = create_report()
    print(f"\nFinal document available at:\n{output_path}")
